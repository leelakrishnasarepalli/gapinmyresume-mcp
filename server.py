#!/usr/bin/env python3
"""Resume Gap Analyzer MCP Server using FastMCP."""

import os
from pathlib import Path
from typing import Annotated, Dict, List, Union

from docx import Document
from dotenv import load_dotenv
from fastmcp import FastMCP
from openai import OpenAI
from pydantic import Field

from prompts import (
    SYSTEM_PROMPT,
    GapAnalysisResult,
    create_analysis_prompt,
)

# Load environment variables
load_dotenv()

# Initialize OpenAI client
openai_client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# Initialize FastMCP server
mcp = FastMCP("resume-gap-analyzer")

# Path to sample resumes directory
SAMPLE_RESUMES_DIR = Path(__file__).parent / "sample-resumes"


def extract_text_from_docx(file_path: Union[str, Path]) -> str:
    """Extract text content from a DOCX file.

    Args:
        file_path: Path to the DOCX file

    Returns:
        Extracted text content
    """
    doc = Document(file_path)
    text_parts = []

    # Extract text from paragraphs
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)

    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_parts.append(cell.text)

    return "\n".join(text_parts)


def get_sample_resumes() -> List[Dict[str, str]]:
    """Get list of sample resume files.

    Returns:
        List of dicts with 'name' and 'path' keys
    """
    if not SAMPLE_RESUMES_DIR.exists():
        return []

    samples = []
    for file_path in SAMPLE_RESUMES_DIR.glob("*.docx"):
        if not file_path.name.startswith("~"):  # Skip temp files
            samples.append({
                "name": file_path.stem,
                "path": str(file_path.absolute()),
            })

    return samples


@mcp.tool()
def analyze_resume_gaps(
    resume_path: Annotated[str, Field(description="Absolute path to resume file (.docx)")],
    job_description: Annotated[str, Field(description="Full text of the job posting")],
) -> GapAnalysisResult:
    """Analyze a resume against a job description to identify gaps and improvements.

    This tool:
    - Extracts text from DOCX resume files
    - Compares resume content against job requirements
    - Identifies missing skills, keywords, and experience gaps
    - Provides prioritized, actionable recommendations
    - Optimizes for both ATS and human readability

    Args:
        resume_path: Absolute file path to the resume document (.docx format)
        job_description: Complete text of the job description to compare against

    Returns:
        Structured analysis with gaps, recommendations, and quick wins

    Raises:
        ValueError: If file format is not supported or file doesn't exist
        Exception: If OpenAI API call fails
    """
    # Validate file exists
    path = Path(resume_path)
    if not path.exists():
        raise ValueError(f"Resume file not found: {resume_path}")

    # Validate file format
    if path.suffix.lower() not in [".docx", ".doc"]:
        raise ValueError("Only .docx files are supported")

    # Extract text from resume
    try:
        resume_text = extract_text_from_docx(path)
    except Exception as e:
        raise ValueError(f"Failed to extract text from resume: {e}")

    if not resume_text.strip():
        raise ValueError("Resume appears to be empty or unreadable")

    # Create analysis prompt
    user_prompt = create_analysis_prompt(resume_text, job_description)

    # Call OpenAI API with structured output
    try:
        completion = openai_client.beta.chat.completions.parse(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": SYSTEM_PROMPT},
                {"role": "user", "content": user_prompt},
            ],
            response_format=GapAnalysisResult,
            temperature=0.7,
        )

        result = completion.choices[0].message.parsed
        if result is None:
            raise Exception("OpenAI returned no parsed result")

        return result

    except Exception as e:
        raise Exception(f"Failed to analyze resume: {e}")


@mcp.resource("resume://sample-{name}")
def get_sample_resume(name: str) -> str:
    """Get the content of a sample resume.

    Args:
        name: Name of the sample resume (without extension)

    Returns:
        Extracted text content of the resume

    Raises:
        ValueError: If sample resume not found
    """
    samples = get_sample_resumes()

    # Find the matching sample
    for sample in samples:
        if sample["name"] == name:
            try:
                return extract_text_from_docx(sample["path"])
            except Exception as e:
                raise ValueError(f"Failed to read sample resume: {e}")

    available = [s["name"] for s in samples]
    raise ValueError(
        f"Sample resume '{name}' not found. Available samples: {', '.join(available)}"
    )


if __name__ == "__main__":
    # Run the MCP server
    mcp.run()
